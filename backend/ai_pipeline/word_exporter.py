import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_word(data, output_dir):
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Extracted Question Paper", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Source: {data['source_file']}")
    doc.add_paragraph(f"Total Pages: {data['total_pages']}")
    doc.add_paragraph(f"Overall Confidence: {data['overall_confidence']}%")
    doc.add_paragraph("")

    for page in data["pages"]:
        # Page heading
        doc.add_heading(f"Page {page['page_number']}", level=1)
        
        # Add extracted content line by line
        content = page["extracted_content"]
        
        available_diagrams = page["diagrams"].copy()
        
        in_table = False
        table_rows = []

        def flush_table():
            nonlocal in_table, table_rows
            if not table_rows:
                return
            # Filter out separator rows like |---|---|
            data_rows = [r for r in table_rows if not re.match(r'^\|[-:| ]+\|$', r)]
            if data_rows:
                # Parse columns
                table_data = []
                for r in data_rows:
                    # Since markdown tables have leading/trailing empty string due to boundary pipes, filter out outer splits, just split and strip.
                    cols = [c.strip() for c in r.strip().strip('|').split('|')]
                    table_data.append(cols)
                
                num_cols = max(len(row) for row in table_data) if table_data else 0
                if num_cols > 0:
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        row_cells = table.rows[i].cells
                        for j, col_data in enumerate(row_data):
                            if j < len(row_cells):
                                row_cells[j].text = col_data
            in_table = False
            table_rows = []
            
        for line in content.split('\n'):
            line = line.strip()
            
            # Table detection
            if line.startswith('|') and line.endswith('|'):
                in_table = True
                table_rows.append(line)
                continue
            else:
                if in_table:
                    flush_table()

            if not line:
                continue

            # Check for DIAGRAM inline
            if re.search(r'\[DIAGRAM:', line):
                if available_diagrams:
                    diagram = available_diagrams.pop(0)
                    if os.path.exists(diagram["image_path"]):
                        doc.add_paragraph("")
                        doc.add_heading("Diagram:", level=3)
                        doc.add_picture(diagram["image_path"], width=Inches(5))
                        doc.add_paragraph(diagram["description"])
                continue

            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            else:
                doc.add_paragraph(line)
        
        # Flush if table ended at the very last line
        if in_table:
            flush_table()
        
        # Add any remaining diagrams not caught by placeholders
        for diagram in available_diagrams:
            if os.path.exists(diagram["image_path"]):
                doc.add_paragraph("")
                doc.add_heading("Diagram:", level=3)
                doc.add_picture(diagram["image_path"], width=Inches(5))
                doc.add_paragraph(diagram["description"])

        doc.add_page_break()

    # Save
    base_output_filename = data["source_file"].replace(".pdf", ".docx")
    output_path = os.path.join(output_dir, base_output_filename)
    
    # Handle duplicate filenames
    if os.path.exists(output_path):
        filename_without_ext, ext = os.path.splitext(base_output_filename)
        counter = 1
        while os.path.exists(output_path):
            output_path = os.path.join(output_dir, f"{filename_without_ext}({counter}){ext}")
            counter += 1
            
    doc.save(output_path)
    print(f"Word file saved to: {output_path}")
    return output_path